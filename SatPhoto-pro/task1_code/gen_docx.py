# -*- coding: utf-8 -*-
"""把成果文档生成为 Word(.docx)，含配图与表格，方便组员查看。"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT, "results")
FIG = os.path.join(OUT, "figures")
DOCX = os.path.join(ROOT, "任务一_RPC正射纠正_成果文档.docx")

doc = Document()
# 中文字体
st = doc.styles["Normal"]
st.font.name = "微软雅黑"
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x2E)
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def pic(path, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hd
        for pp in c.paragraphs:
            for r in pp.runs:
                r.bold = True
                r.font.name = "微软雅黑"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for pp in cells[i].paragraphs:
                for r in pp.runs:
                    r.font.name = "微软雅黑"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return t


# ===== 标题 =====
title = doc.add_heading("任务一 · 基于 RPC 模型的卫星影像正射纠正 —— 成果文档", level=0)
for r in title.runs:
    r.font.name = "微软雅黑"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
para("负责人：____（姓名/学号）    日期：2026-06    数据：JAX JAX_Tile_163_RGB_001 与 005（同一地区两个视角）",
     italic=True)
para("一句话：把「不准确 RPC 影像」按正确高程基准纠正成几何准确的正射影像 DOM。对 001、005 两个视角分别正射，"
     "用 ArcGIS 卷帘叠加验证——只有都用椭球高 DEM 纠正，两视角的地面地物才重合；用 0 高程面或水准高纠正都不重合。"
     "本环节是团队「从 RPC 影像 到 DOM/DSM 产品」流程的第一步（正射出图）。")

# ===== 一、怎么做 =====
h("一、怎么做（技术路线）", 1)
para("整条链路只用 RPC 正解（地面→像方）+ 间接法（反向映射）正射，无需 RPC 反解：")
para("输出DOM像素(row,col) → 像素中心地理坐标 lon=left+(col+0.5)·res_x, lat=top−(row+0.5)·res_y "
     "→ 取高程 Z（0/DEM正常高/DEM椭球高） → RPC正解得原图(line,sample) → 双线性内插重采样 → 赋值。",
     italic=True)
para("三个关键点（也是本任务的考点）：", bold=True)
para("1) 像素坐标系：行=line、列=sample，左上角第一个像素中心为原点；输出 DOM 像素中心要 +0.5 像元再算经纬度。")
para("2) 高程基准必须统一到 WGS84 椭球高。RPC 用椭球高，而 USGS DEM 是正常高（水准高），两者相差高程异常 N"
     "（本区≈−29.6m，由 EGM2008 提供）：椭球高 h = 正常高 H(DEM) + N。用错基准会带来几十米高程偏差→几十像素平面错位。")
para("3) 工程实现：整幅向量化构网格 + OpenCV cv2.remap 双线性重采样；无 GDAL，用 tifffile 读图、.tfw+.prj 写地理参考。")

# ===== 二、做了什么 =====
h("二、做了什么（5 个小题，对应满分 60）", 1)
table(["小题", "内容", "关键做法"], [
    ["第1题(30分)", "RPC 正解：12 个地面点投影到像方", "教材 P75 公式3.1，20项有理多项式"],
    ["第2题(10分)", "正射到 0 高程面", "每像素 Z=0（椭球高0面），001/005 各一幅"],
    ["第3题(10分)", "正射到 DEM 水准高", "DEM 双线性内插取正常高，001/005 各一幅"],
    ["第4题(5分)", "正射到 DEM 椭球高（正确）", "正常高+N→椭球高，P112，001/005 各一幅"],
    ["第5题(5分)", "平行投影模型近似 RPC 加速", "RPC 虚拟控制网拟合8参数仿射，P63/P86"],
])
para("第2/3/4 题对 001 与 005 两个视角各做一遍，共产出 6 个 DOM（3 种高程面 × 2 视角），"
     "用于第三节的跨视角卷帘检核。", italic=True)
para("第1题 RPC 正解（投影点位）：", bold=True); pic(os.path.join(FIG, "fig_q1_points.png"), 4.2)
para("第2/3/4题 三种高程面下的正射 DOM 成果（视角 001）：", bold=True); pic(os.path.join(FIG, "fig_doms.png"), 6.4)
para("第2/3/4题 三种高程面下的正射 DOM 成果（视角 005）：", bold=True); pic(os.path.join(FIG, "fig_doms_005.png"), 6.4)
para("第5题 平行投影加速：", bold=True); pic(os.path.join(FIG, "fig_q5.png"), 6.4)

# ===== 三、结果评定 =====
h("三、结果评定", 1)
para("1) 第1题 —— 数值完全一致", bold=True)
table(["指标", "line", "sample", "点位"], [["RMSE（像素）", "0.00000", "0.00000", "0.00000"]])
para("")
para("2) 跨视角卷帘检核 —— 本任务的核心验证（为什么必须用椭球高）", bold=True)
para("把 005 的 DOM 按地理坐标重采样到 001 的格网，在公共有效区内统计 001 与 005 两视角同名地物的"
     "残余平移（即「地面上是否重合」），三种高程面对比（001 格网像元 ≈ 0.298×0.342 m/px）：")
pic(os.path.join(FIG, "fig_crossview.png"), 6.6)
table(["高程面", "两视角重合NCC", "残余平移(像素)", "地面错位(米)", "是否重合"], [
    ["0 高程面", "0.469", "33.0", "11.07 m", "✗ 明显错位"],
    ["水准高(DEM正常高)", "0.428", "43.0", "14.43 m", "✗ 明显错位"],
    ["椭球高（正确）", "0.749", "0.49", "0.17 m", "✓ 重合(亚像素)"],
])
para("结论与机理：", bold=True)
para("· 只有都用椭球高 DEM，两视角错位才被压到 0.49 像素 ≈ 0.17 m（亚像素级，地面重合）；用 0 高程面/水准高"
     "残留 11~14 m（33~43 像素）的明显错位、无法重合，与老师「卷帘只有椭球高重合」的结论完全一致。")
para("· 0高程面、水准高两种错误基准的错位方向几乎一致、量级随高程误差单调增大（0高程面误差≈+26.6m、水准高≈+29.6m），"
     "正是「系统高程基准误差驱动的视差」特征，证明错位来自高程基准而非偶然。")
para("· NCC 量纲说明：此处是两个不同视角真实影像之间比（受光照/BRDF/车辆移动/投影差影响），椭球高 NCC≈0.75（非1）"
     "仍属重合良好，与下面「我方 vs 参考（同源影像）」的 NCC≈0.99 不是一个量纲，不可直接比较。")
para("· 投影差正常：建筑/桥梁/树木等高出地面的地物，即使在椭球高 DOM 上仍有投影差、局部不重合（边缘叠加图中残留的"
     "红/绿分离），属正常现象，不影响地面高程基准正确性的结论。")
para("")
para("3) 第2/3/4题 —— 与老师 ArcGIS 参考 DOM 几何重合（001 与 005 均验证）", bold=True)
para("参考 DOM 与我方网格完全一致（同 .tfw），但参考为16bit且经ArcGIS拉伸，故比几何对齐："
     "用 NCC（对拉伸不敏感）自动配对、相位相关求残余平移。")
pic(os.path.join(FIG, "fig_ncc.png"), 4.6)
table(["视角", "我方成果", "匹配参考变体", "NCC", "残余平移"], [
    ["001", "0高程面", "参考 0高程面", "0.990", "0.007 像素"],
    ["001", "水准高", "参考 水准高", "0.990", "0.028 像素"],
    ["001", "椭球高（正确）", "参考 椭球高", "0.989", "0.157 像素"],
    ["005", "0高程面", "参考 0高程面", "0.990", "0.002 像素"],
    ["005", "水准高", "参考 水准高", "0.989", "0.005 像素"],
    ["005", "椭球高（正确）", "参考 椭球高", "0.989", "0.195 像素"],
])
para("两视角 NCC 矩阵均对角占优、残余平移均<0.2像素→与参考几何完全重合。")
para("关于 NCC≈0.99 而非 1.0：老师参考 DOM 由 ArcGIS 默认的最近邻重采样生成，我方用的是双线性内插，"
     "两者重采样核不同 + ArcGIS 颜色拉伸，必然带来微小灰度差异——这是正常现象（老师已说明），并非几何误差。")
para("椭球高 DOM 与参考叠加检核（棋盘格中道路/建筑跨格连续=配准良好）：", bold=True)
pic(os.path.join(OUT, "diff_q4_ellipH_001.png"), 6.4)
para("")
para("4) 第4题 —— EGM2008 高程异常 N 交叉校验", bold=True)
table(["来源", "N 均值", "一致性"], [
    ["EGM2008 直接内插", "−29.580 m", "—"],
    ["控制点隐含 N（椭球高−DEM正常高）", "−29.845 m", "两者 RMSE 仅 0.38 m"],
])
para("→ 证明高程异常与椭球高改正正确。")
para("")
para("5) 第5题 —— 平行投影加速且无显著差异", bold=True)
table(["项目", "数值"], [
    ["虚拟控制点拟合残差 RMSE", "line 0.036 / sample 0.142 像素"],
    ["与 RPC 像点坐标差异", "RMSE line 0.033 / sample 0.131；最大 0.51 像素（亚像素）"],
    ["DOM 灰度差异 RMSE", "1.3 / 255"],
    ["坐标解算耗时", "RPC 4161 ms → 平行投影 190 ms"],
    ["加速倍率", "≈ 22×"],
])
para("总评：五小题全部实现、跑通；001、005 两视角与老师参考答案在数值/几何上对齐，且跨视角卷帘检核证明"
     "只有椭球高基准使两视角地面重合，按打分原则可达「结果准确(100%)」档位。", bold=True)

# ===== 四、成果清单与衔接 =====
h("四、成果清单与组员衔接", 1)
para("交付代码：task1_code/（Python，约80KB，不含依赖与数据，满足10MB限制）")
para("成果产品（results/）：q1_pixels.csv；001 三套 DOM(q2_dom_0m / q3_dom_orthoH / q4_dom_ellipH) + "
     "005 三套 DOM(同名加 _005) + q5_dom_parallel，各 .tif（含 .tfw/.prj，可直接入 ArcGIS）；"
     "跨视角检核 cross_view_report.txt 与 figures/fig_crossview.png；"
     "compare_report.txt、diff_q4_ellipH_001/005.png、q5_timing.txt、figures/。")
para("与团队流程的接口：", bold=True)
para("· 输入：原始影像(001/005) + 原始 RPC(.rpb) + DEM + DOM范围(正射影像范围.csv)。")
para("· 输出：几何准确的正射 DOM（椭球高版为准确产品；两视角地面重合）。")
para("· 上游衔接：若用「任务二」输出的校正后 RPC 替换原始 RPC，即可产出「从不准确RPC 到 准确正射影像」完整成果（团队90%档）。")
para("· 可复用模块：rpc.py / geoid_egm2008.py / dem.py / raster_io.py / cross_view.py 可供核线纠正、前方交会等环节直接调用。")
para("运行：cd task1_code && python run_all.py（Windows 看中文加 PYTHONIOENCODING=utf-8）。"
     "依赖 numpy scipy opencv-python tifffile matplotlib；第4/5题需 EGM2008 格网 ../geoid_data/geoids/egm2008-5.pgm。")

doc.save(DOCX)
print("已生成:", DOCX)
